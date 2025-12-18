"""
Doküman oluşturucu modülü - python-docx ile DOCX, docx2pdf ile PDF oluşturma.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert


class DocumentGenerator:
    """DOCX ve PDF oluşturucu sınıfı."""
    
    def __init__(self, db=None, footer_info: dict = None, pdf_settings: dict = None):
        """Doküman oluşturucuyu başlat.
        
        Args:
            db: Veritabanı nesnesi (ayarları okumak için)
            footer_info: Altbilgi bilgileri {"phone": "", "website": "", "instagram": ""}
            pdf_settings: PDF ayarları {"font": "", "title_size": 18, "subtitle_size": 14, "content_size": 11, "time_size": 10}
        """
        self.db = db
        
        # Veritabanından ayarları oku (eğer db varsa ve manuel parametre yoksa)
        if db and not footer_info:
            self.footer_info = {
                "phone": db.get_setting("footer_phone", ""),
                "website": db.get_setting("footer_website", ""),
                "instagram": db.get_setting("footer_instagram", "")
            }
        else:
            self.footer_info = footer_info or {}
        
        if db and not pdf_settings:
            self.pdf_settings = {
                "font": db.get_setting("pdf_font", "Comic Sans MS"),
                "title_size": db.get_setting("pdf_title_size", "18"),
                "subtitle_size": db.get_setting("pdf_subtitle_size", "14"),
                "content_size": db.get_setting("pdf_content_size", "11"),
                "time_size": db.get_setting("pdf_time_size", "10")
            }
        else:
            self.pdf_settings = pdf_settings or {}
        
        # PDF ayarlarını değişkenlere ata
        self.font_name = self.pdf_settings.get("font", "Comic Sans MS")
        self.title_size = int(self.pdf_settings.get("title_size", 18))
        self.subtitle_size = int(self.pdf_settings.get("subtitle_size", 14))
        self.content_size = int(self.pdf_settings.get("content_size", 11))
        self.time_size = int(self.pdf_settings.get("time_size", 10))
    
    def _get_meal_style(self, meal_type: str):
        """Öğün türüne göre renk ve görünen adı döndür."""
        main_meals = {
            "kahvalti": ("KAHVALTI", RGBColor(231, 76, 60)),  # Kırmızı
            "ogle": ("ÖĞLE YEMEĞİ", RGBColor(231, 76, 60)),   # Kırmızı
            "aksam": ("AKŞAM", RGBColor(231, 76, 60))         # Kırmızı
        }
        snacks = ["ara_ogun_1", "ara_ogun_2", "ara_ogun_3"]
        drinks = ["ozel_icecek"]
        
        if meal_type in main_meals:
            return main_meals[meal_type]
        elif meal_type in snacks:
            return ("ARA ÖĞÜN", RGBColor(39, 174, 96))  # Yeşil
        elif meal_type in drinks:
            return ("ARA ÖĞÜN", RGBColor(243, 156, 18))  # Sarı/Turuncu
        else:
            return ("ÖĞÜN", RGBColor(39, 174, 96))
    
    def _add_footer(self, doc):
        """Altbilgi ekle."""
        footer_parts = []
        if self.footer_info.get("phone"):
            footer_parts.append(self.footer_info['phone'])
        if self.footer_info.get("website"):
            footer_parts.append(self.footer_info['website'])
        if self.footer_info.get("instagram"):
            footer_parts.append(f"@{self.footer_info['instagram']}")
        
        if footer_parts:
            footer_text = "     |     ".join(footer_parts)
            
            # Her section için footer ekle
            for section in doc.sections:
                footer = section.footer
                footer.is_linked_to_previous = False
                p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(footer_text)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(127, 140, 141)  # Gri
                run.font.name = self.font_name
    
    def create_diet_document(self, file_path: str, diet_program: list, 
                             template_name: str, pool_type: str, bki_group: str):
        """Diyet programı DOCX ve PDF oluştur.
        
        Args:
            file_path: Dosya yolu (uzantısız veya .docx'li)
            diet_program: Diyet programı verisi
            template_name: Kalıp adı
            pool_type: Havuz türü
            bki_group: BKİ grubu
            
        Returns:
            tuple: (docx_path, pdf_path)
        """
        # Dosya yollarını hazırla
        base_path = file_path.replace('.pdf', '').replace('.docx', '')
        docx_path = f"{base_path}.docx"
        pdf_path = f"{base_path}.pdf"
        
        # Yeni doküman oluştur
        doc = Document()
        
        # Varsayılan font ve satır aralığı ayarla
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.content_size)
        # Satır aralığını 1.15 yap
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        
        # Her gün için içerik oluştur
        for i, day_data in enumerate(diet_program):
            day_num = day_data["day"]
            meals = day_data["meals"]
            
            # Gün başlığı (ortalı, yeşil, kalın)
            day_title = doc.add_paragraph()
            day_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = day_title.add_run(f"{day_num}. Gün")
            run.bold = True
            run.font.size = Pt(self.title_size)
            run.font.color.rgb = RGBColor(39, 174, 96)  # Yeşil
            run.font.name = self.font_name
            
            # Su hatırlatma (ortalı, yeşil, italik)
            water = doc.add_paragraph()
            water.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = water.add_run("(Her gün 2,5 litre su içmeyi unutmayın)")
            run.italic = True
            run.font.size = Pt(self.subtitle_size - 2)
            run.font.color.rgb = RGBColor(39, 174, 96)  # Yeşil
            run.font.name = self.font_name
            
            # Sabah aç karnına metin
            morning = doc.add_paragraph()
            run = morning.add_run("Sabah aç karnına 1 bardak su")
            run.bold = True
            run.font.size = Pt(self.content_size)
            run.font.name = self.font_name
            
            # Öğünler
            for meal in meals:
                time = meal["time"]
                meal_type = meal.get("meal_type", "")
                recipe_text = meal["recipe_text"]
                
                # Öğün başlığını ve rengini al
                display_name, color = self._get_meal_style(meal_type)
                
                # Öğün başlığı (kalın, renkli)
                meal_para = doc.add_paragraph()
                meal_para.paragraph_format.space_before = Pt(12)  # Öğünler arası boşluk
                run = meal_para.add_run(f"{display_name}:{time}")
                run.bold = True
                run.font.size = Pt(self.subtitle_size - 1)
                run.font.color.rgb = color
                run.font.name = self.font_name
                
                # İçerik - virgülle ayrılmış öğeleri liste olarak göster (girintili)
                items = [item.strip() for item in recipe_text.split(",")]
                for item in items:
                    if item:
                        # Manuel bullet ile girintili paragraf
                        item_para = doc.add_paragraph()
                        item_para.paragraph_format.left_indent = Cm(1.5)
                        item_para.paragraph_format.first_line_indent = Cm(-0.5)
                        run = item_para.add_run(f"• {item}")
                        run.font.size = Pt(self.content_size)
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Siyah
                        run.font.name = self.font_name
            
            # Son gün değilse sayfa sonu ekle
            if i < len(diet_program) - 1:
                doc.add_page_break()
        
        # Footer ekle
        self._add_footer(doc)
        
        # DOCX kaydet
        doc.save(docx_path)
        
        # PDF'e dönüştür
        try:
            convert(docx_path, pdf_path)
        except Exception as e:
            print(f"PDF dönüştürme hatası: {e}")
            pdf_path = None
        
        return docx_path, pdf_path
    
    def create_program(self, save_path: str, diet_data: dict, template_id: int,
                       start_date, bki_group: str, pool_type: str, excluded_foods: str = "") -> bool:
        """Diyet programı oluştur (diet_creator.py ile uyumlu).
        
        Args:
            save_path: Kayıt yolu
            diet_data: Diyet verileri
            template_id: Şablon ID
            start_date: Başlangıç tarihi (QDate)
            bki_group: BKİ grubu
            pool_type: Havuz türü
            excluded_foods: Hariç tutulacak yiyecekler
            
        Returns:
            bool: Başarılı ise True
        """
        import random
        from datetime import timedelta
        
        try:
            # Şablonu al
            template = self.db.get_template(template_id)
            if not template:
                return False
            
            # Gün sayısı
            days = diet_data.get("days", 4)
            
            # Hariç tutulacak kelimeler
            exclude_words = [w.strip().lower() for w in excluded_foods.split(",") if w.strip()]
            
            # Her gün için program oluştur
            diet_program = []
            
            for day in range(1, days + 1):
                day_meals = []
                
                # Şablondaki öğünleri işle
                template_meals = template.get("meals", [])
                
                for meal in template_meals:
                    meal_type = meal.get("meal_type", "")
                    meal_time = meal.get("time", "")
                    
                    # Bu öğün türü için tarifleri al
                    recipes = self.db.get_all_recipes(pool_type, meal_type)
                    
                    # Hariç tutulanları filtrele
                    if exclude_words:
                        filtered_recipes = []
                        for r in recipes:
                            recipe_text = r.get(f"bki_{bki_group}", r.get("bki_21_25", "")).lower()
                            if not any(word in recipe_text for word in exclude_words):
                                filtered_recipes.append(r)
                        recipes = filtered_recipes if filtered_recipes else recipes
                    
                    # Rastgele tarif seç
                    if recipes:
                        selected = random.choice(recipes)
                        recipe_text = selected.get(f"bki_{bki_group}", selected.get("bki_21_25", ""))
                    else:
                        recipe_text = "Tarif bulunamadı"
                    
                    day_meals.append({
                        "time": meal_time,
                        "meal_type": meal_type,
                        "recipe_text": recipe_text
                    })
                
                diet_program.append({
                    "day": day,
                    "meals": day_meals
                })
            
            # Dokümanı oluştur
            self.create_diet_document(
                file_path=save_path,
                diet_program=diet_program,
                template_name=template.get("name", ""),
                pool_type=pool_type,
                bki_group=bki_group
            )
            
            return True
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return False


# Geriye uyumluluk için alias
PDFGenerator = DocumentGenerator
