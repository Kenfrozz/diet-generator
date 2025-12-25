from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import datetime

class DOCXGenerator:
    def __init__(self, footer_info=None):
        self.footer_info = footer_info or {}
        
        # Renkler (PDF ile aynı)
        self.colors = {
            'green': RGBColor(0x27, 0xae, 0x60),      # #27ae60
            'red': RGBColor(0xe7, 0x4c, 0x3c),        # #e74c3c
            'orange': RGBColor(0xf3, 0x9c, 0x12),     # #f39c12
            'gray': RGBColor(0x7f, 0x8c, 0x8d),       # #7f8c8d
            'dark': RGBColor(0x2c, 0x3e, 0x50),       # #2c3e50
        }

    def _set_cell_shading(self, cell, hex_color):
        """Hücre arka plan rengini ayarla."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def create_diet_docx(self, file_path, diet_program, patient_name, start_date, 
                         template_name, bki_group, excluded_foods, combination_code,
                         patient_info=None):
        """
        Diyet programı DOCX oluştur (PDF ile aynı stil).
        
        Args:
            patient_info: Opsiyonel hasta bilgileri dict:
                - weight: Başlangıç kilosu
                - height: Boy (cm)
                - birth_year: Doğum yılı
                - end_date: Liste bitiş tarihi (kontrol tarihi)
        """
        doc = Document()
        
        # --- Sayfa Ayarları ---
        section = doc.sections[0]
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2.5)
        
        # --- Font Ayarları (Comic Sans benzeri) ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Comic Sans MS'
        font.size = Pt(11)
        
        # Heading stilleri güncelle
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'Comic Sans MS'
            heading_style.font.color.rgb = self.colors['green']
        
        # === KAPAK SAYFASI ===
        if patient_info:
            self._create_cover_page(doc, patient_name, start_date, patient_info)
            doc.add_page_break()
        
        # === DİYET PROGRAMI ===
        for i, day_data in enumerate(diet_program):
            self._create_day_page(doc, day_data)
            
            # Son gün değilse sayfa sonu
            if i < len(diet_program) - 1:
                doc.add_page_break()

        # --- Footer ---
        self._add_footer(doc)
        
        doc.save(file_path)
        return file_path
    
    def _create_cover_page(self, doc, patient_name, start_date, patient_info):
        """Kapak sayfası oluştur - Kullanıcı formatı."""
        
        # Hasta bilgileri
        weight = patient_info.get('weight', 0)
        height = patient_info.get('height', 0)
        birth_year = patient_info.get('birth_year', 0)
        end_date = patient_info.get('end_date', '')
        
        # Hesaplamalar
        current_year = datetime.datetime.now().year
        yas = current_year - birth_year if birth_year else 0
        
        # BKİ hesapla
        height_m = height / 100 if height else 1
        bki = weight / (height_m * height_m) if height_m else 0
        
        # İdeal ve geçmemesi gereken kilo hesapla
        if yas < 35:
            ideal_kilo = height_m * height_m * 21
            gecmemesi_gereken = height_m * height_m * 27
        elif 35 <= yas <= 45:
            ideal_kilo = height_m * height_m * 22
            gecmemesi_gereken = height_m * height_m * 28
        else:
            ideal_kilo = height_m * height_m * 23
            gecmemesi_gereken = height_m * height_m * 30
        
        # === AD SOYAD ===
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run("Ad Soyad: ")
        run.font.size = Pt(14)
        run.font.color.rgb = self.colors['gray']
        run = name_para.add_run(patient_name.upper())
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = self.colors['dark']
        
        doc.add_paragraph()
        
        # === BAŞLANGIÇ KİLOSU - BOY - YAŞ (Tablo ile iki yana yaslı) ===
        info_table = doc.add_table(rows=1, cols=3)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.autofit = True
        
        # Sütun genişlikleri
        for cell in info_table.columns[0].cells:
            cell.width = Cm(5.5)
        for cell in info_table.columns[1].cells:
            cell.width = Cm(5.5)
        for cell in info_table.columns[2].cells:
            cell.width = Cm(5.5)
        
        row = info_table.rows[0]
        
        # Sol hücre - Başlangıç
        if weight:
            left_cell = row.cells[0]
            left_para = left_cell.paragraphs[0]
            left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = left_para.add_run("Başlangıç: ")
            run.font.size = Pt(12)
            run.font.color.rgb = self.colors['gray']
            run = left_para.add_run(f"{weight:.1f} kg")
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = self.colors['dark']
        
        # Orta hücre - Boy
        if height:
            mid_cell = row.cells[1]
            mid_para = mid_cell.paragraphs[0]
            mid_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = mid_para.add_run("Boy: ")
            run.font.size = Pt(12)
            run.font.color.rgb = self.colors['gray']
            run = mid_para.add_run(f"{height:.0f} cm")
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = self.colors['dark']
        
        # Sağ hücre - Yaş
        if yas:
            right_cell = row.cells[2]
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = right_para.add_run("Yaş: ")
            run.font.size = Pt(12)
            run.font.color.rgb = self.colors['gray']
            run = right_para.add_run(f"{yas}")
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = self.colors['dark']
        
        doc.add_paragraph()
        
        # === KONTROL TARİHİ ===
        if end_date:
            control_para = doc.add_paragraph()
            control_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = control_para.add_run("Kontrol Tarihi: ")
            run.font.size = Pt(12)
            run.font.color.rgb = self.colors['gray']
            run = control_para.add_run(end_date.upper())
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = self.colors['green']
        
        for _ in range(2):
            doc.add_paragraph()
        
        # === BKİ ===
        if bki:
            bki_para = doc.add_paragraph()
            bki_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = bki_para.add_run("BKİ: ")
            run.font.size = Pt(14)
            run.font.color.rgb = self.colors['gray']
            run = bki_para.add_run(f"{bki:.1f}")
            run.bold = True
            run.font.size = Pt(18)
            if bki < 26:
                run.font.color.rgb = self.colors['green']
            elif bki < 30:
                run.font.color.rgb = self.colors['orange']
            else:
                run.font.color.rgb = self.colors['red']
        
        # === İDEAL KİLO ===
        if ideal_kilo:
            ideal_para = doc.add_paragraph()
            ideal_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = ideal_para.add_run("İdeal Kilonuz: ")
            run.font.size = Pt(13)
            run.font.color.rgb = self.colors['gray']
            run = ideal_para.add_run(f"{ideal_kilo:.1f} kg")
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = self.colors['green']
        
        # === GEÇMEMESİ GEREKEN KİLO ===
        if gecmemesi_gereken:
            max_para = doc.add_paragraph()
            max_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = max_para.add_run("Geçmemeniz Gereken Kilo: ")
            run.font.size = Pt(13)
            run.font.color.rgb = self.colors['gray']
            run = max_para.add_run(f"{gecmemesi_gereken:.1f} kg")
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = self.colors['red']

    def _create_day_page(self, doc, day_data):
        """Gün sayfası oluştur (PDF stili ile)."""
        day_num = day_data["day"]
        meals = day_data["meals"]
        
        # Gün başlığı (yeşil, ortalı)
        day_title = doc.add_paragraph()
        day_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = day_title.add_run(f"{day_num}. Gün")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = self.colors['green']
        
        # Su hatırlatma (yeşil, italik, ortalı)
        water = doc.add_paragraph()
        water.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = water.add_run("(Her gün 2,5 litre su içmeyi unutmayın)")
        run.italic = True
        run.font.size = Pt(12)
        run.font.color.rgb = self.colors['green']
        
        # Sabah aç karnına
        morning = doc.add_paragraph()
        run = morning.add_run("Sabah aç karnına 1 bardak su")
        run.bold = True
        run.font.size = Pt(11)
        
        doc.add_paragraph()  # Boşluk
        
        # Öğünler
        for meal in meals:
            self._add_meal(doc, meal)
    
    def _add_meal(self, doc, meal):
        """Öğün ekle (PDF stili ile)."""
        time = meal["time"]
        meal_name = meal["meal_name"]
        meal_type = meal.get("meal_type", "")
        recipe_text = meal["recipe_text"]
        
        # Öğün başlığı
        meal_para = doc.add_paragraph()
        
        # Öğün tipine göre stil
        main_meals = {"kahvalti": "KAHVALTI", "ogle": "ÖĞLE YEMEĞİ", "aksam": "AKŞAM"}
        snacks = ["ara_ogun_1", "ara_ogun_2", "ara_ogun_3"]
        drinks = ["ozel_icecek"]
        
        if meal_type in main_meals:
            display_name = f"{main_meals[meal_type]}: {time}"
            color = self.colors['red']
        elif meal_type in snacks:
            display_name = f"ARA ÖĞÜN: {time}"
            color = self.colors['green']
        elif meal_type in drinks:
            display_name = f"ÖZEL İÇECEK: {time}"
            color = self.colors['orange']
        else:
            display_name = f"{meal_name.upper()}: {time}"
            color = self.colors['green']
        
        run = meal_para.add_run(display_name)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = color
        
        # İçerik - bullet points (virgülle ayrılmış)
        items = [item.strip() for item in recipe_text.split(",")]
        for item in items:
            if item:
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.paragraph_format.left_indent = Cm(1)
                run = bullet.add_run(item)
                run.font.size = Pt(11)

    def _add_footer(self, doc):
        """Footer ekle."""
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_parts = []
        if self.footer_info.get("phone"):
            footer_parts.append(self.footer_info['phone'])
        if self.footer_info.get("website"):
            footer_parts.append(self.footer_info['website'])
        if self.footer_info.get("instagram"):
            footer_parts.append(f"@{self.footer_info['instagram']}")
        
        if footer_parts:
            run = p.add_run("     |     ".join(footer_parts))
            run.font.size = Pt(9)
            run.font.color.rgb = self.colors['gray']
